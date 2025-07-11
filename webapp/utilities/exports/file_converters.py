# Copyright (C) 2025 David West Brown

# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at

#     http://www.apache.org/licenses/LICENSE-2.0

# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

# File export and conversion utilities

from docx import Document
from docx.shared import RGBColor
from docx.opc.part import Part
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import pandas as pd
import polars as pl
import zipfile
from typing import Union

from webapp.utilities.corpus import get_corpus_data_manager


def convert_to_excel(df: Union[pl.DataFrame, pd.DataFrame]) -> bytes:
    """
    Convert a DataFrame to an Excel file in memory.

    Parameters
    ----------
    df : Union[pl.DataFrame, pd.DataFrame]
        The DataFrame to be converted to Excel format.

    Returns
    -------
    bytes
        The Excel file as a bytes object.
    """
    output = BytesIO()
    writer = pd.ExcelWriter(output, engine='xlsxwriter')

    # Convert to pandas only if needed
    if hasattr(df, 'to_pandas'):
        pandas_df = df.to_pandas()
    else:
        pandas_df = df

    pandas_df.to_excel(writer, index=False, header=True)
    writer.close()
    processed_data = output.getvalue()

    # Clean up memory
    output.close()
    del pandas_df

    return processed_data


def add_alt_chunk(doc: Document, html: str) -> None:
    """
    Add an HTML altChunk to a Word document.

    Parameters
    ----------
    doc : Document
        The Word document to which the altChunk will be added.
    html : str
        The HTML string to embed as an altChunk.

    Returns
    -------
    None
    """
    package = doc.part.package
    partname = package.next_partname('/word/altChunk%d.html')
    alt_part = Part(
        partname,
        'text/html',
        html.encode(),
        package
        )
    r_id = doc.part.relate_to(
        alt_part,
        RELATIONSHIP_TYPE.A_F_CHUNK
        )
    alt_chunk = OxmlElement('w:altChunk')
    alt_chunk.set(qn('r:id'), r_id)
    doc.element.body.sectPr.addprevious(alt_chunk)


def convert_to_word(
        html_string: str,
        tag_html: str,
        doc_key: str,
        ) -> bytes:
    """
    Convert HTML content and tag counts into a Word document and
    return it as bytes.

    Parameters
    ----------
    html_string : str
        The HTML string containing the document content and styles.
    tag_html : str
        The HTML string representing highlighted tags to embed in the document.
    doc_key : str
        The document key or title to use in the Word file.
    tag_counts : pd.DataFrame
        A DataFrame containing tag frequency information to include as a table.

    Returns
    -------
    bytes
        The generated Word document as a bytes object.
    """
    doc_html = html_string.split('</style>')
    style_sheet_str = doc_html[0] + '</style>'
    html_str = doc_html[1]
    doc_html = ('<!DOCTYPE html><html><head>' + style_sheet_str +
                '</head><body>' + tag_html +
                '<br><br>' + html_str +
                '</body></html>')
    download_file = Document()
    title = download_file.add_heading(doc_key)
    title.style.font.color.rgb = RGBColor(0, 0, 0)
    heading = download_file.add_heading('Highlighted tags:', 3)
    heading.style.font.color.rgb = RGBColor(0, 0, 0)

    # add html
    add_alt_chunk(download_file, doc_html)
    output = BytesIO()
    download_file.save(output)
    processed_data = output.getvalue()
    return processed_data


def convert_corpus_to_zip(
        session_id: str,
        corpus_type: str,
        file_type="parquet"
        ) -> bytes:
    """
    Convert all tables in a corpus to a ZIP archive of Parquet or CSV files.

    Parameters
    ----------
    session_id : str
        The session ID for which the corpus is to be exported.
    corpus_type : str
        The type of corpus to export (e.g., 'target' or 'reference').
    file_type : str, optional
        The file format for export: 'parquet' (default) or 'csv'.

    Returns
    -------
    bytes
        The ZIP archive as a bytes object, with each table as a file.
    """
    # Use the new data manager system
    manager = get_corpus_data_manager(session_id, corpus_type)

    if not manager.is_ready():
        raise ValueError(f"Corpus data not ready for {corpus_type}")

    zip_buf = BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as file_zip:
        # Get all available data keys from the manager
        available_keys = manager.get_available_keys()

        for table_name in available_keys:
            # Get the data for this key
            df = manager.get_data(table_name)

            # Skip None values (data not generated yet)
            if df is None:
                continue

            # Ensure we have a valid DataFrame
            if not hasattr(df, 'to_pandas'):
                continue

            try:
                if file_type == "parquet":
                    parquet_data = df.to_pandas().to_parquet()
                    file_zip.writestr(f"{table_name}.parquet", parquet_data)
                else:
                    csv_data = df.to_pandas().to_csv()
                    file_zip.writestr(f"{table_name}.csv", csv_data)
            except Exception:
                # Skip tables that can't be converted, but continue with others
                continue

    processed_data = zip_buf.getvalue()
    return processed_data


def convert_to_zip(
        tok_pl: pl.DataFrame,
        tagset: str
        ) -> bytes:
    """
    Convert tokenized corpus data to a ZIP archive of tagged text files.

    Parameters
    ----------
    tok_pl : pl.DataFrame
        The Polars DataFrame containing tokenized corpus data.
    tagset : str
        The tagset to use for tagging ('pos' or 'ds').

    Returns
    -------
    bytes
        The ZIP archive as a bytes object, with each document
        as a tagged text file.
    """
    zip_buf = BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as file_zip:
        for id in tok_pl.get_column("doc_id").unique().to_list():
            if tagset == "pos":
                df = (
                    tok_pl
                    .filter(pl.col("doc_id") == id)
                    .group_by(
                        ["pos_id", "pos_tag"], maintain_order=True
                        )
                    .agg(pl.col("token").str.concat(""))
                    .with_columns(
                        pl.col("token").str.strip_chars()
                        )
                    .with_columns(
                        pl.col("token").str.replace_all(" ", "_")
                        )
                    .with_columns(
                        pl.when(pl.col("pos_tag") == "Y")
                        .then(pl.col("pos_tag").str.replace(
                            "Y", "", literal=True
                            ))
                        .when(pl.col("pos_tag") == "FU")
                        .then(pl.col("pos_tag").str.replace(
                            "FU", "", literal=True
                            ))
                        .otherwise(pl.col("pos_tag")))
                    .with_columns(
                        pl.concat_str(
                            pl.col("token"), pl.lit("|"), pl.col("pos_tag")
                            )
                        )
                    .with_columns(
                        pl.col("token").str.replace_all(r"\|$", "")
                        )
                    )
            else:
                df = (
                    tok_pl
                    .filter(pl.col("doc_id") == id)
                    .group_by(["ds_id", "ds_tag"], maintain_order=True)
                    .agg(pl.col("token").str.concat(""))
                    .with_columns(
                        pl.col("token").str.strip_chars()
                        )
                    .with_columns(
                        pl.col("token").str.replace_all(" ", "_")
                        )
                    .with_columns(
                        pl.when(
                            pl.col("ds_tag") == "Untagged"
                            )
                        .then(
                            pl.col("ds_tag").str.replace(
                                "Untagged",
                                "",
                                literal=True)
                            )
                        .otherwise(pl.col("ds_tag")))
                    .with_columns(
                        pl.concat_str(
                            pl.col("token"),
                            pl.lit("|"),
                            pl.col("ds_tag")
                            )
                        )
                    .with_columns(
                        pl.col("token").str.replace_all(r"\|$", "")
                        )
                    )
            doc = " ".join(df.get_column("token").to_list())
            file_zip.writestr(id + "_tagged" + ".txt", doc)
    processed_data = zip_buf.getvalue()

    return processed_data
